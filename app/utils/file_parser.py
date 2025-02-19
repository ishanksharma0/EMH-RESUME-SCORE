# app/utils/file_parser.py

from io import BytesIO
import logging
from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image
import re
from zipfile import ZipFile
import xml.etree.ElementTree as ET
import win32com.client
import tempfile

logger = logging.getLogger(__name__)

def parse_pdf_or_docx(file_buffer: BytesIO, filename: str) -> str:
    """
    Determines the file type (PDF, DOC, DOCX, or image) and extracts text accordingly.
    :param file_buffer: File buffer of the uploaded file.
    :param filename: Name of the uploaded file.
    :return: Extracted text content as a string.
    """
    try:
        if filename.lower().endswith(".pdf"):
            return parse_pdf(file_buffer)
        elif filename.lower().endswith(".docx"):
            return parse_docx(file_buffer)
        elif filename.lower().endswith(".doc"):
            return parse_doc(file_buffer)
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            return image_to_text(file_buffer)  # Handle image to text conversion
        else:
            raise ValueError("Unsupported file format. Only PDF, DOCX, DOC, and image formats are supported.")
    except Exception as e:
        logger.error(f"Error parsing file '{filename}': {str(e)}", exc_info=True)
        raise

def parse_pdf(file_buffer: BytesIO) -> str:
    """
    Extracts text from a PDF file, including hyperlinks.
    :param file_buffer: File buffer of the uploaded PDF file.
    :return: Extracted text content as a string, including hyperlinks.
    """
    try:
        logger.info("Parsing PDF file")
        reader = PdfReader(file_buffer)
        text = ""
        hyperlinks = []

        # Extract text from each page and gather hyperlinks from metadata if available
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text

            # Extract hyperlinks from annotations (if available)
            if "/Annots" in page:
                annotations = page["/Annots"]
                for annotation in annotations:
                    # Check if annotation is an IndirectObject and resolve it properly
                    if isinstance(annotation, dict):
                        if "/A" in annotation and "/URI" in annotation["/A"]:
                            hyperlinks.append(annotation["/A"]["/URI"])

        # Join the hyperlinks into a single string (one per line)
        hyperlinks_text = '\n'.join(hyperlinks)
        return text.strip() + '\n' + hyperlinks_text

    except Exception as e:
        logger.error(f"Error reading PDF file: {str(e)}", exc_info=True)
        raise

def parse_docx(file_buffer: BytesIO) -> str:
    """
    Extracts text from a DOCX file, including hyperlinks and headers/footers.
    :param file_buffer: File buffer of the uploaded DOCX file.
    :return: Extracted text content as a string, including hyperlinks.
    """
    try:
        logger.info("Parsing DOCX file")
        doc = Document(file_buffer)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'

        # Extract hyperlinks from the document (href="...") in the HTML of the document
        hyperlinks = extract_hyperlinks_from_docx(file_buffer)
        
        # Extract header and footer text
        header_footer_text = extract_header_footer(doc)

        return text.strip() + '\n' + hyperlinks + '\n' + header_footer_text

    except Exception as e:
        logger.error(f"Error reading DOCX file: {str(e)}", exc_info=True)
        raise

def extract_hyperlinks_from_docx(file_buffer: BytesIO) -> str:
    """
    Extracts hyperlinks from a DOCX file by scanning for <a> tags in the document's HTML.
    :param file_buffer: The file buffer of the DOCX file.
    :return: A string containing all hyperlinks found in the document.
    """
    try:
        # Load the DOCX file as a zip and read the XML content
        docx = ZipFile(file_buffer)
        hyperlinks = []
        for file in docx.namelist():
            if "hyperlink" in file:
                content = docx.read(file)
                links = re.findall(r'href="([^"]+)"', content.decode('utf-8'))
                hyperlinks.extend(links)
        return '\n'.join(hyperlinks)

    except Exception as e:
        logger.error(f"Error extracting hyperlinks from DOCX file: {str(e)}", exc_info=True)
        raise

def extract_header_footer(doc) -> str:
    """
    Extract text from headers and footers in a DOCX file.
    :param doc: The Document object for DOCX file.
    :return: Text from the headers and footers.
    """
    text = ""
    # Extract text from headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            text += paragraph.text + '\n'
        
        # Extract text from footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            text += paragraph.text + '\n'

    return text.strip()

def parse_doc(file_buffer: BytesIO) -> str:
    """
    Extracts text from a DOC file using COM (pywin32).
    :param file_buffer: File buffer of the uploaded DOC file.
    :return: Extracted text content as a string.
    """
    try:
        logger.info("Parsing DOC file")
        # Create a temporary file to save the DOC file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(file_buffer.read())
            temp_filename = temp_file.name
        
        # Initialize COM client for Word
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(temp_filename)

        # Extract text from the DOC file
        doc_text = doc.Content.Text
        
        # Close Word document
        doc.Close()
        word.Quit()

        return doc_text.strip()
    
    except Exception as e:
        logger.error(f"Error reading DOC file: {str(e)}", exc_info=True)
        raise

def image_to_text(file_buffer: BytesIO) -> str:
    """
    Extract text from an image using Tesseract OCR.
    :param file_buffer: The image file buffer.
    :return: Extracted text content as a string.
    """
    try:
        logger.info("Extracting text from image")
        image = Image.open(file_buffer)
        text = pytesseract.image_to_string(image)
        return text.strip()
    
    except Exception as e:
        logger.error(f"Error processing image file: {str(e)}", exc_info=True)
        raise
